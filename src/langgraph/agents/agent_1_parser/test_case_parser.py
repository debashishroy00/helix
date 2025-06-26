"""
Helix Automation: Agent 1 - Test Case Parser
LangGraph Node for parsing manual test cases from multiple formats
"""

import asyncio
import re
import io
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import time

from src.langgraph.state.helix_state import HelixAutomationState


async def agent_1_parser(state: HelixAutomationState) -> HelixAutomationState:
    """
    LangGraph Node: Parse manual test cases from any format
    Updates shared state with parsed results
    """
    
    print(f"🤖 Agent 1 (Parser): Processing input...")
    start_time = time.time()
    
    try:
        raw_input = state["raw_input"]
        input_metadata = state.get("input_metadata", {})
        
        # Detect input format if not provided
        input_format = input_metadata.get("format") or detect_input_format(raw_input)
        
        # Parse based on format
        if input_format == "excel":
            parsed_result = await parse_excel_test_case(raw_input, input_metadata)
        elif input_format == "word":
            parsed_result = await parse_word_document(raw_input, input_metadata)
        elif input_format == "gherkin":
            parsed_result = parse_gherkin_feature(raw_input)
        elif input_format == "json":
            parsed_result = parse_json_test_case(raw_input)
        elif input_format == "plain_text":
            parsed_result = parse_plain_text(raw_input)
        elif input_format == "csv":
            parsed_result = parse_csv_test_case(raw_input)
        else:
            # Fallback to plain text parsing
            parsed_result = parse_plain_text(raw_input)
            parsed_result["warnings"] = [f"Unknown format '{input_format}', defaulted to plain text parsing"]
        
        # Update state
        state["parsed_test_case"] = parsed_result["test_case"]
        state["parsing_confidence"] = parsed_result["confidence"]
        state["parsing_errors"] = parsed_result.get("errors", [])
        state["current_agent"] = "Agent 1 (Parser)"
        
        # Update performance metrics
        execution_time = time.time() - start_time
        state["performance_metrics"]["agent_1_duration"] = execution_time
        
        # Update input metadata with detected format
        state["input_metadata"]["detected_format"] = input_format
        state["input_metadata"]["parsing_method"] = parsed_result.get("method", "unknown")
        
        step_count = len(parsed_result["test_case"].get("steps", []))
        print(f"✅ Parsed {step_count} test steps from {input_format} format ({parsed_result['confidence']:.1%} confidence) in {execution_time:.2f}s")
        
        if parsed_result.get("warnings"):
            for warning in parsed_result["warnings"]:
                print(f"⚠️  {warning}")
        
    except Exception as e:
        state["parsing_errors"] = [str(e)]
        state["parsing_confidence"] = 0.0
        print(f"❌ Parsing failed: {e}")
    
    return state


def detect_input_format(raw_input: Union[str, bytes]) -> str:
    """
    Detect the format of the input test case
    """
    
    if isinstance(raw_input, bytes):
        # Binary data - check file signatures
        if raw_input.startswith(b'PK'):  # ZIP-based formats (Excel, Word)
            if b'xl/' in raw_input:
                return "excel"
            elif b'word/' in raw_input:
                return "word"
        elif raw_input.startswith(b'%PDF'):
            return "pdf"
        else:
            # Try to decode as text
            try:
                raw_input = raw_input.decode('utf-8')
            except UnicodeDecodeError:
                return "binary"
    
    if isinstance(raw_input, str):
        raw_lower = raw_input.lower().strip()
        
        # JSON format detection
        if (raw_lower.startswith('{') and raw_lower.endswith('}')) or \
           (raw_lower.startswith('[') and raw_lower.endswith(']')):
            try:
                import json
                json.loads(raw_input)
                return "json"
            except:
                pass
        
        # Gherkin format detection
        if any(keyword in raw_lower for keyword in ['feature:', 'scenario:', 'given ', 'when ', 'then ', 'and ']):
            return "gherkin"
        
        # CSV format detection
        if ',' in raw_input and '\n' in raw_input:
            lines = raw_input.split('\n')
            if len(lines) > 1:
                # Check if multiple lines have similar comma counts
                comma_counts = [line.count(',') for line in lines[:5] if line.strip()]
                if len(set(comma_counts)) <= 2 and max(comma_counts) > 0:
                    return "csv"
        
        # Test case pattern detection
        step_patterns = [
            r'step\s*\d+', r'step\s*\w+', r'\d+\.\s+', r'\d+\)\s+',
            r'test\s*case', r'test\s*step', r'action:', r'expected:'
        ]
        
        if any(re.search(pattern, raw_lower) for pattern in step_patterns):
            return "plain_text"
        
        # Default to plain text for string input
        return "plain_text"
    
    return "unknown"


def parse_plain_text(raw_input: str) -> Dict[str, Any]:
    """
    Parse plain text test cases with various formats
    """
    
    lines = [line.strip() for line in raw_input.strip().split('\n') if line.strip()]
    
    # Extract test case metadata
    test_case = {
        "test_id": extract_test_id(raw_input),
        "title": extract_test_title(lines),
        "description": extract_test_description(lines),
        "steps": []
    }
    
    # Parse test steps
    steps = parse_test_steps_from_lines(lines)
    test_case["steps"] = steps
    
    # Calculate confidence based on structure quality
    confidence = calculate_parsing_confidence(test_case, raw_input)
    
    return {
        "test_case": test_case,
        "confidence": confidence,
        "errors": [],
        "method": "plain_text_parser"
    }


def parse_gherkin_feature(raw_input: str) -> Dict[str, Any]:
    """
    Parse Gherkin/BDD format test cases
    """
    
    lines = [line.strip() for line in raw_input.strip().split('\n')]
    
    test_case = {
        "test_id": "GHERKIN_AUTO",
        "title": "Gherkin Feature",
        "description": "",
        "steps": []
    }
    
    current_step = None
    step_number = 1
    
    for line in lines:
        line_lower = line.lower()
        
        # Extract feature title
        if line_lower.startswith('feature:'):
            test_case["title"] = line[8:].strip()
        
        # Extract scenario
        elif line_lower.startswith('scenario:'):
            test_case["description"] = line[9:].strip()
        
        # Parse Gherkin steps
        elif any(line_lower.startswith(keyword) for keyword in ['given ', 'when ', 'then ', 'and ', 'but ']):
            if current_step:
                test_case["steps"].append(current_step)
            
            # Determine action type
            if line_lower.startswith('given '):
                action_type = "setup"
                description = line[6:].strip()
            elif line_lower.startswith('when '):
                action_type = "action"
                description = line[5:].strip()
            elif line_lower.startswith('then '):
                action_type = "verification"
                description = line[5:].strip()
            elif line_lower.startswith('and '):
                action_type = "continuation"
                description = line[4:].strip()
            elif line_lower.startswith('but '):
                action_type = "exception"
                description = line[4:].strip()
            
            current_step = {
                "step_number": step_number,
                "action": action_type,
                "description": description,
                "expected_result": "",
                "test_data": {},
                "gherkin_keyword": line.split()[0].lower()
            }
            step_number += 1
    
    # Add the last step
    if current_step:
        test_case["steps"].append(current_step)
    
    confidence = 0.9 if test_case["steps"] else 0.3
    
    return {
        "test_case": test_case,
        "confidence": confidence,
        "errors": [],
        "method": "gherkin_parser"
    }


def parse_json_test_case(raw_input: str) -> Dict[str, Any]:
    """
    Parse JSON format test cases
    """
    
    try:
        import json
        data = json.loads(raw_input)
        
        # Handle different JSON structures
        if isinstance(data, dict):
            if "test_case" in data:
                test_case = data["test_case"]
            else:
                test_case = data
        elif isinstance(data, list):
            # Array of steps
            test_case = {
                "test_id": "JSON_AUTO",
                "title": "JSON Test Case",
                "description": "Imported from JSON",
                "steps": []
            }
            
            for i, step_data in enumerate(data, 1):
                if isinstance(step_data, dict):
                    step = {
                        "step_number": i,
                        "action": step_data.get("action", "unknown"),
                        "description": step_data.get("description", step_data.get("step", "")),
                        "expected_result": step_data.get("expected", step_data.get("expected_result", "")),
                        "test_data": step_data.get("data", step_data.get("test_data", {}))
                    }
                    test_case["steps"].append(step)
        
        # Ensure required fields exist
        test_case.setdefault("test_id", "JSON_AUTO")
        test_case.setdefault("title", "JSON Test Case")
        test_case.setdefault("description", "")
        test_case.setdefault("steps", [])
        
        confidence = 0.95 if test_case["steps"] else 0.5
        
        return {
            "test_case": test_case,
            "confidence": confidence,
            "errors": [],
            "method": "json_parser"
        }
        
    except json.JSONDecodeError as e:
        return {
            "test_case": {"test_id": "ERROR", "title": "JSON Parse Error", "description": "", "steps": []},
            "confidence": 0.0,
            "errors": [f"JSON parsing failed: {str(e)}"],
            "method": "json_parser"
        }


def parse_csv_test_case(raw_input: str) -> Dict[str, Any]:
    """
    Parse CSV format test cases
    """
    
    lines = raw_input.strip().split('\n')
    if not lines:
        return {
            "test_case": {"test_id": "CSV_EMPTY", "title": "Empty CSV", "description": "", "steps": []},
            "confidence": 0.0,
            "errors": ["Empty CSV input"],
            "method": "csv_parser"
        }
    
    # Parse CSV header
    header = [col.strip().strip('"') for col in lines[0].split(',')]
    
    test_case = {
        "test_id": "CSV_AUTO",
        "title": "CSV Test Case",
        "description": "",
        "steps": []
    }
    
    # Map common column names
    column_map = {}
    for i, col in enumerate(header):
        col_lower = col.lower()
        if any(keyword in col_lower for keyword in ['step', 'action', 'description']):
            column_map['description'] = i
        elif any(keyword in col_lower for keyword in ['expected', 'result']):
            column_map['expected_result'] = i
        elif any(keyword in col_lower for keyword in ['data', 'input']):
            column_map['test_data'] = i
    
    # Parse data rows
    for row_num, line in enumerate(lines[1:], 1):
        if not line.strip():
            continue
            
        values = [val.strip().strip('"') for val in line.split(',')]
        
        step = {
            "step_number": row_num,
            "action": extract_action_from_step(values[0] if values else ""),
            "description": values[column_map.get('description', 0)] if values else "",
            "expected_result": values[column_map.get('expected_result', 1)] if len(values) > 1 else "",
            "test_data": parse_test_data_string(values[column_map.get('test_data', 2)]) if len(values) > 2 else {}
        }
        test_case["steps"].append(step)
    
    confidence = 0.8 if test_case["steps"] else 0.2
    
    return {
        "test_case": test_case,
        "confidence": confidence,
        "errors": [],
        "method": "csv_parser"
    }


async def parse_excel_test_case(raw_input: bytes, input_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """
    Parse Excel file containing test cases
    Note: This requires pandas and openpyxl to be installed
    """
    
    try:
        import pandas as pd
        from io import BytesIO
        
        # Read Excel file
        excel_file = BytesIO(raw_input)
        
        # Try to read the first sheet
        try:
            df = pd.read_excel(excel_file, sheet_name=0)
        except Exception as e:
            return {
                "test_case": {"test_id": "EXCEL_ERROR", "title": "Excel Parse Error", "description": "", "steps": []},
                "confidence": 0.0,
                "errors": [f"Excel reading failed: {str(e)}"],
                "method": "excel_parser"
            }
        
        # Extract test case information from first row or dedicated cells
        test_case = {
            "test_id": find_value_in_dataframe(df, ["test id", "test_id", "id"]) or "EXCEL_AUTO",
            "title": find_value_in_dataframe(df, ["test title", "title", "test name"]) or "Excel Test Case",
            "description": find_value_in_dataframe(df, ["description", "desc", "summary"]) or "",
            "steps": []
        }
        
        # Find step columns
        step_column = find_column_in_dataframe(df, ["step", "action", "description"])
        expected_column = find_column_in_dataframe(df, ["expected", "result", "expected result"])
        data_column = find_column_in_dataframe(df, ["data", "test data", "input"])
        
        # Parse test steps
        step_number = 1
        for index, row in df.iterrows():
            step_desc = str(row[step_column]) if step_column and pd.notna(row[step_column]) else ""
            
            if step_desc and step_desc.lower() not in ['nan', 'none', '']:
                step = {
                    "step_number": step_number,
                    "action": extract_action_from_step(step_desc),
                    "description": step_desc,
                    "expected_result": str(row[expected_column]) if expected_column and pd.notna(row[expected_column]) else "",
                    "test_data": parse_test_data_string(str(row[data_column])) if data_column and pd.notna(row[data_column]) else {}
                }
                test_case["steps"].append(step)
                step_number += 1
        
        confidence = 0.9 if test_case["steps"] else 0.3
        
        return {
            "test_case": test_case,
            "confidence": confidence,
            "errors": [],
            "method": "excel_parser"
        }
        
    except ImportError:
        return {
            "test_case": {"test_id": "EXCEL_ERROR", "title": "Excel Parse Error", "description": "", "steps": []},
            "confidence": 0.0,
            "errors": ["pandas and openpyxl required for Excel parsing"],
            "method": "excel_parser"
        }


async def parse_word_document(raw_input: bytes, input_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """
    Parse Word document containing test cases
    Note: This requires python-docx to be installed
    """
    
    try:
        from docx import Document
        from io import BytesIO
        
        # Read Word document
        doc_file = BytesIO(raw_input)
        doc = Document(doc_file)
        
        # Extract text from document
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        # Parse as plain text
        plain_text = '\n'.join(full_text)
        result = parse_plain_text(plain_text)
        result["method"] = "word_parser"
        
        return result
        
    except ImportError:
        return {
            "test_case": {"test_id": "WORD_ERROR", "title": "Word Parse Error", "description": "", "steps": []},
            "confidence": 0.0,
            "errors": ["python-docx required for Word document parsing"],
            "method": "word_parser"
        }
    except Exception as e:
        return {
            "test_case": {"test_id": "WORD_ERROR", "title": "Word Parse Error", "description": "", "steps": []},
            "confidence": 0.0,
            "errors": [f"Word document parsing failed: {str(e)}"],
            "method": "word_parser"
        }


# Helper functions

def extract_test_id(text: str) -> str:
    """Extract test ID from text"""
    
    patterns = [
        r'test\s*id\s*:?\s*([^\n\r]+)',
        r'tc\s*-?\s*(\d+)',
        r'test\s*case\s*:?\s*([^\n\r]+)',
        r'id\s*:?\s*([^\n\r]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return "AUTO_GENERATED"


def extract_test_title(lines: List[str]) -> str:
    """Extract test title from lines"""
    
    for line in lines[:5]:  # Check first 5 lines
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in ['test case', 'title', 'test:']):
            # Extract title part
            for separator in [':', '-', '|']:
                if separator in line:
                    return line.split(separator, 1)[1].strip()
            return line.strip()
    
    # Default to first non-empty line
    return lines[0] if lines else "Untitled Test Case"


def extract_test_description(lines: List[str]) -> str:
    """Extract test description from lines"""
    
    description_lines = []
    in_description = False
    
    for line in lines:
        line_lower = line.lower()
        
        if any(keyword in line_lower for keyword in ['description:', 'desc:', 'summary:']):
            in_description = True
            # Add content after the keyword
            for separator in [':', '-', '|']:
                if separator in line:
                    desc_part = line.split(separator, 1)[1].strip()
                    if desc_part:
                        description_lines.append(desc_part)
                    break
        elif in_description and not any(keyword in line_lower for keyword in ['step', 'test', 'action']):
            description_lines.append(line)
        elif any(keyword in line_lower for keyword in ['step', 'test', 'action']) and description_lines:
            break
    
    return ' '.join(description_lines).strip()


def parse_test_steps_from_lines(lines: List[str]) -> List[Dict[str, Any]]:
    """Parse test steps from text lines"""
    
    steps = []
    current_step = None
    step_number = 1
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        
        # Check if this line starts a new step
        if is_step_line(line_clean):
            # Save previous step
            if current_step:
                steps.append(current_step)
            
            # Start new step
            current_step = {
                "step_number": step_number,
                "action": extract_action_from_step(line_clean),
                "description": clean_step_description(line_clean),
                "expected_result": "",
                "test_data": {}
            }
            step_number += 1
        
        elif current_step and is_expected_result_line(line_clean):
            current_step["expected_result"] = extract_expected_result(line_clean)
        
        elif current_step and is_test_data_line(line_clean):
            current_step["test_data"] = parse_test_data_string(line_clean)
        
        elif current_step:
            # Continuation of current step description
            current_step["description"] += " " + line_clean
    
    # Add the last step
    if current_step:
        steps.append(current_step)
    
    return steps


def is_step_line(line: str) -> bool:
    """Check if line represents a test step"""
    
    line_lower = line.lower()
    
    # Numbered steps
    if re.match(r'^\d+[\.\)]\s+', line):
        return True
    
    # Step keywords
    step_keywords = ['step', 'action', 'click', 'enter', 'navigate', 'verify', 'check', 'select', 'fill']
    if any(line_lower.startswith(keyword) for keyword in step_keywords):
        return True
    
    # Common test patterns
    if re.match(r'^(given|when|then|and)\s+', line_lower):
        return True
    
    return False


def is_expected_result_line(line: str) -> bool:
    """Check if line represents expected result"""
    
    line_lower = line.lower()
    keywords = ['expected', 'result', 'should', 'verify', 'assert']
    return any(keyword in line_lower for keyword in keywords)


def is_test_data_line(line: str) -> bool:
    """Check if line contains test data"""
    
    line_lower = line.lower()
    keywords = ['data:', 'input:', 'username:', 'password:', 'value:', 'text:']
    return any(keyword in line_lower for keyword in keywords)


def extract_action_from_step(step_text: str) -> str:
    """Extract action type from step description"""
    
    step_lower = step_text.lower()
    
    # Navigation actions
    if any(keyword in step_lower for keyword in ['navigate', 'go to', 'open', 'visit']):
        return "navigate"
    
    # Click actions
    elif any(keyword in step_lower for keyword in ['click', 'press', 'tap']):
        return "click"
    
    # Input actions
    elif any(keyword in step_lower for keyword in ['enter', 'type', 'input', 'fill']):
        return "input"
    
    # Selection actions
    elif any(keyword in step_lower for keyword in ['select', 'choose', 'pick']):
        return "select"
    
    # Verification actions
    elif any(keyword in step_lower for keyword in ['verify', 'check', 'assert', 'confirm']):
        return "verify"
    
    # Wait actions
    elif any(keyword in step_lower for keyword in ['wait', 'pause', 'delay']):
        return "wait"
    
    return "action"


def clean_step_description(step_text: str) -> str:
    """Clean step description by removing step numbering and keywords"""
    
    # Remove step numbering
    cleaned = re.sub(r'^\d+[\.\)]\s*', '', step_text)
    
    # Remove common prefixes
    prefixes = ['step:', 'action:', 'given ', 'when ', 'then ', 'and ']
    for prefix in prefixes:
        if cleaned.lower().startswith(prefix):
            cleaned = cleaned[len(prefix):].strip()
            break
    
    return cleaned.strip()


def extract_expected_result(line: str) -> str:
    """Extract expected result from line"""
    
    # Remove expected result keywords
    keywords = ['expected:', 'result:', 'should:', 'verify:', 'assert:']
    line_lower = line.lower()
    
    for keyword in keywords:
        if keyword in line_lower:
            return line[line_lower.index(keyword) + len(keyword):].strip()
    
    return line.strip()


def parse_test_data_string(data_str: str) -> Dict[str, Any]:
    """Parse test data from string"""
    
    if not data_str or data_str.lower() in ['nan', 'none', '']:
        return {}
    
    data = {}
    
    # Try JSON format first
    try:
        import json
        return json.loads(data_str)
    except:
        pass
    
    # Parse key-value pairs
    # Handle formats like "username=test@example.com, password=secret"
    if '=' in data_str:
        pairs = data_str.split(',')
        for pair in pairs:
            if '=' in pair:
                key, value = pair.split('=', 1)
                data[key.strip()] = value.strip()
    
    # Handle colon-separated format
    elif ':' in data_str:
        pairs = data_str.split(',')
        for pair in pairs:
            if ':' in pair:
                key, value = pair.split(':', 1)
                data[key.strip()] = value.strip()
    
    # Single value
    else:
        data['value'] = data_str.strip()
    
    return data


def find_value_in_dataframe(df, keywords: List[str]) -> Optional[str]:
    """Find value in dataframe by searching for keywords"""
    
    for keyword in keywords:
        # Search in column names
        for col in df.columns:
            if keyword.lower() in str(col).lower():
                # Return first non-null value in this column
                for value in df[col]:
                    if pd.notna(value) and str(value).lower() not in ['nan', 'none', '']:
                        return str(value)
        
        # Search in cell values
        for col in df.columns:
            for value in df[col]:
                if pd.notna(value) and keyword.lower() in str(value).lower():
                    return str(value)
    
    return None


def find_column_in_dataframe(df, keywords: List[str]) -> Optional[str]:
    """Find column name in dataframe by searching for keywords"""
    
    for keyword in keywords:
        for col in df.columns:
            if keyword.lower() in str(col).lower():
                return col
    
    return None


def calculate_parsing_confidence(test_case: Dict[str, Any], raw_input: str) -> float:
    """Calculate confidence score for parsing quality"""
    
    confidence = 0.0
    
    # Base confidence for having steps
    if test_case.get("steps"):
        confidence += 0.4
        
        # Bonus for multiple steps
        step_count = len(test_case["steps"])
        confidence += min(0.2, step_count * 0.05)
    
    # Bonus for having test metadata
    if test_case.get("title") and test_case["title"] != "Untitled Test Case":
        confidence += 0.1
    
    if test_case.get("description"):
        confidence += 0.1
    
    if test_case.get("test_id") and test_case["test_id"] != "AUTO_GENERATED":
        confidence += 0.1
    
    # Bonus for step quality
    for step in test_case.get("steps", []):
        if step.get("expected_result"):
            confidence += 0.02
        if step.get("test_data"):
            confidence += 0.02
        if step.get("action") != "action":  # Non-generic action
            confidence += 0.01
    
    # Bonus for structured input
    if any(keyword in raw_input.lower() for keyword in ['step', 'test case', 'expected', 'action']):
        confidence += 0.1
    
    return min(1.0, confidence)


# Test function for Agent 1
async def test_agent_1():
    """Test Agent 1 in isolation"""
    
    from src.langgraph.state.helix_state import create_initial_state
    
    # Test cases for different formats
    test_cases = [
        {
            "name": "Plain Text Test",
            "input": """
            Test Case: Login to Salesforce
            Description: Test user login functionality
            
            Step 1: Navigate to https://login.salesforce.com
            Step 2: Enter username: test@company.com
            Step 3: Enter password: password123
            Step 4: Click Login button
            Expected: User should be logged in successfully
            """,
            "metadata": {"format": "plain_text"}
        },
        {
            "name": "Gherkin Test",
            "input": """
            Feature: Salesforce Login
            Scenario: User logs into Salesforce
            Given I am on the Salesforce login page
            When I enter username "test@company.com"
            And I enter password "password123"
            And I click the Login button
            Then I should be logged in successfully
            """,
            "metadata": {"format": "gherkin"}
        },
        {
            "name": "CSV Test",
            "input": """Step,Expected Result,Test Data
Navigate to login page,Login page displayed,
Enter username,Username field filled,test@company.com
Enter password,Password field filled,password123
Click login,User logged in,""",
            "metadata": {"format": "csv"}
        }
    ]
    
    print("🧪 Testing Agent 1 (Test Case Parser)...")
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n📋 Test {i}: {test_case['name']}")
        print("-" * 40)
        
        # Create initial state
        initial_state = create_initial_state(
            raw_input=test_case["input"],
            input_metadata=test_case["metadata"]
        )
        
        # Run Agent 1
        result_state = await agent_1_parser(initial_state)
        
        # Display results
        parsed_case = result_state.get("parsed_test_case", {})
        confidence = result_state.get("parsing_confidence", 0)
        errors = result_state.get("parsing_errors", [])
        
        print(f"✅ Parsing Results:")
        print(f"   - Confidence: {confidence:.1%}")
        print(f"   - Steps parsed: {len(parsed_case.get('steps', []))}")
        print(f"   - Test ID: {parsed_case.get('test_id', 'N/A')}")
        print(f"   - Title: {parsed_case.get('title', 'N/A')}")
        
        if errors:
            print(f"❌ Errors:")
            for error in errors:
                print(f"     {error}")
        
        # Show first step as example
        if parsed_case.get("steps"):
            first_step = parsed_case["steps"][0]
            print(f"📝 First step example:")
            print(f"     Action: {first_step.get('action', 'N/A')}")
            print(f"     Description: {first_step.get('description', 'N/A')[:60]}...")
    
    print(f"\n🎯 Agent 1 testing complete!")


if __name__ == "__main__":
    asyncio.run(test_agent_1())